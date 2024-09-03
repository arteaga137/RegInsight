from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os
import json
import pandas as pd

def create_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    r_style = OxmlElement('w:rStyle')
    r_style.set(qn('w:val'), 'Hyperlink')
    rPr.append(r_style)
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_word_report_with_analysis(summaries, filename=None):
    today = datetime.today().strftime('%Y-%m-%d')
    
    if filename is None:
        filename = f"regulatory_updates_report_{today}.docx"
    
    storage_dir = os.path.expanduser("~/Documents/Regulatory Reports")
    os.makedirs(storage_dir, exist_ok=True)
    
    dated_folder = os.path.join(storage_dir, today)
    os.makedirs(dated_folder, exist_ok=True)
    
    full_path = os.path.join(dated_folder, filename)
    
    base_filename, file_extension = os.path.splitext(filename)
    counter = 1
    while os.path.exists(full_path):
        new_filename = f"{base_filename}_{counter}{file_extension}"
        full_path = os.path.join(dated_folder, new_filename)
        counter += 1
    
    doc = Document()
    doc.add_heading(f"Regulatory Updates Report {today}", 0)

    for summary in summaries:
        try:
            data = json.loads(summary)
        except Exception as e:
            print(f"Error decoding JSON content: {e}")
            continue

        title_section = data.get("Title of the Regulation", {})
        if title_section:
            english_title = title_section.get("English", "N/A")
            doc.add_heading(english_title, level=1)
            for language in data.get("Title of the Regulation", {}):
                doc.add_paragraph(f"{language}: {title_section[language]}", style='List Bullet')

        tags = data.get("Tags", [])
        if tags:
            doc.add_heading("Tags", level=2)
            for tag in tags:
                doc.add_paragraph(tag, style='List Bullet')
        
        doc.add_heading("From", level=2)
        doc.add_paragraph(data.get("From", "N/A"))
        
        doc.add_heading("Subject", level=2)
        doc.add_paragraph(data.get("Subject", "N/A"))
        
        doc.add_heading("Jurisdiction", level=2)
        doc.add_paragraph(data.get("Jurisdiction", "N/A"))
        
        doc.add_heading("Date", level=2)
        doc.add_paragraph(data.get("Date", "N/A"))
        
        doc.add_heading("Subject Matter", level=2)
        doc.add_paragraph(data.get("Subject Matter", "N/A"))

        doc.add_heading("Summary", level=2)
        doc.add_paragraph(data.get("Summary", "N/A"))

        analysis_section = data.get("Analysis", {})
        if analysis_section:
            doc.add_heading("Analysis", level=2)
            doc.add_paragraph(analysis_section.get("Purpose and Scope", "N/A"))

        doc.add_heading("List of Relevant Regulations or Legal Measures", level=2)
        for regulation in data.get("Relevant Regulations or Legal Measures", []):
            doc.add_paragraph(regulation, style='List Bullet')

        business_impact_section = data.get("Business Impact", {})
        if business_impact_section:
            doc.add_heading("Business Impact", level=2)
            doc.add_paragraph(business_impact_section.get("Overview", "N/A"))

        conclusion_section = data.get("Conclusion", {})
        if conclusion_section:
            doc.add_heading("Conclusion", level=2)
            doc.add_paragraph(conclusion_section.get("Overview", "N/A"))

        doc.add_heading("Relevant Links", level=2)
        for link in data.get("Relevant Links", []):
            description = link.get("description", "N/A")
            url = link.get("url", "#")
            paragraph = doc.add_paragraph(style='List Bullet')
            create_hyperlink(paragraph, url, description)

        doc.add_heading("Note", level=2)
        doc.add_paragraph(data.get("Note", "N/A"))
        
        doc.add_page_break()

    doc.save(full_path)
    print(f"Report saved as {full_path}")

def create_excel_report_with_analysis(summaries, filename=None):
    today = datetime.today().strftime('%Y-%m-%d')
    
    if filename is None:
        filename = f"regulatory_updates_report_{today}.xlsx"
    
    downloads_dir = os.path.expanduser("~/Documents/Regulatory Reports")
    
    dated_folder = os.path.join(downloads_dir, today)
    os.makedirs(dated_folder, exist_ok=True)
    
    full_path = os.path.join(dated_folder, filename)
    
    base_filename, file_extension = os.path.splitext(filename)
    counter = 1
    while os.path.exists(full_path):
        new_filename = f"{base_filename}_{counter}{file_extension}"
        full_path = os.path.join(dated_folder, new_filename)
        counter += 1
    
    rows = []
    for summary in summaries:
        try:
            data = json.loads(summary)
            title_section = data.get("Title of the Regulation", {})
            english_title = title_section.get("English", "N/A")
            original_title = title_section.get("Original Language", "N/A")
            tags = ", ".join(data.get("Tags", []))
            from_ = data.get("From", "N/A")
            subject = data.get("Subject", "N/A")
            jurisdiction = data.get("Jurisdiction", "N/A")
            date = data.get("Date", "N/A")
            subject_matter = data.get("Subject Matter", "N/A")
            summary_text = data.get("Summary", "N/A")
            analysis = data.get("Analysis", {}).get("Purpose and Scope", "N/A")
            relevant_regs = "; ".join(data.get("Relevant Regulations or Legal Measures", []))
            business_impact = data.get("Business Impact", {}).get("Overview", "N/A")
            conclusion = data.get("Conclusion", {}).get("Overview", "N/A")
            relevant_links = "; ".join([f'{link.get("description", "N/A")}: {link.get("url", "#")}' for link in data.get("Relevant Links", [])])
            note = data.get("Note", "N/A")
            
            rows.append([
                english_title, original_title, tags, from_, subject, 
                jurisdiction, date, subject_matter, summary_text, analysis, 
                relevant_regs, business_impact, conclusion, relevant_links, note
            ])
        except Exception as e:
            print(f"Error processing summary for Excel: {e}")
    
    df = pd.DataFrame(rows, columns=[
        "English Title", "Original Title", "Tags", "From", "Subject", 
        "Jurisdiction", "Date", "Subject Matter", "Summary", "Analysis", 
        "Relevant Regulations or Legal Measures", "Business Impact", 
        "Conclusion", "Relevant Links", "Note"
    ])
    
    df.to_excel(full_path, index=False)
    print(f"Excel report saved as {full_path}")